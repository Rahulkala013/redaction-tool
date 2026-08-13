import io
import re
import spacy
from faker import Faker
from docx import Document

# Initialize Faker and spaCy
fake = Faker('en_IN')
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    import subprocess
    import sys
    subprocess.run([sys.executable, "-m", "spacy", "download", "en_core_web_sm"])
    nlp = spacy.load("en_core_web_sm")

# Dictionary to maintain consistent fake replacements across the document
replacement_map = {}
redact_cache = {}

# --- REGEX PATTERNS ---
PATTERNS = {
    "EMAIL": r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,7}\b',
    "PHONE": r'(\+?\d{1,3}[-.\s]?)?(\(?\d{2,5}\)?[-.\s]?)?\d{3,4}[-.\s]?\d{4}\b',
    "IP_ADDRESS": r'\b(?:\d{1,3}\.){3}\d{1,3}\b',
    "CREDIT_CARD": r'\b(?:\d[ -]*?){13,16}\b',
    "SSN_PAN": r'\b[A-Z]{5}[0-9]{4}[A-Z]{1}\b|\b\d{3}-\d{2}-\d{4}\b', 
    "AADHAAR": r'\b\d{4}\s\d{4}\s\d{4}\b',
    "DATE_OF_BIRTH": r'\b(0[1-9]|[12][0-9]|3[01])[- /.](0[1-9]|1[012])[- /.](19|20)\d\d\b'
}

def get_fake_value(entity_text, entity_type):
    """Returns a consistent fake value for a given entity text."""
    if entity_text in replacement_map:
        return replacement_map[entity_text]
    
    if entity_type == "PERSON":
        fake_val = fake.name()
    elif entity_type == "EMAIL":
        fake_val = fake.email()
    elif entity_type == "PHONE":
        fake_val = fake.phone_number()
    elif entity_type == "ORG":
        fake_val = fake.company()
    elif entity_type == "GPE" or entity_type == "LOC": 
        fake_val = fake.address().replace('\n', ', ')
    elif entity_type == "SSN_PAN":
        fake_val = fake.bothify(text='?????####?').upper() 
    elif entity_type == "AADHAAR":
        fake_val = fake.bothify(text='#### #### ####')
    elif entity_type == "CREDIT_CARD":
        fake_val = fake.credit_card_number()
    elif entity_type == "DATE_OF_BIRTH":
        fake_val = fake.date_of_birth(minimum_age=18, maximum_age=90).strftime('%d/%m/%Y')
    elif entity_type == "IP_ADDRESS":
        fake_val = fake.ipv4()
    else:
        fake_val = "[REDACTED]"

    replacement_map[entity_text] = fake_val
    return fake_val

def redact_text_regex(text):
    """Applies regex-based redaction patterns to a string."""
    redacted_text = text
    for label, pattern in PATTERNS.items():
        matches = re.finditer(pattern, redacted_text)
        for match in reversed(list(matches)):
            original = match.group(0)
            if "Order" in text[max(0, match.start()-10):match.start()] or "Ticket" in text[max(0, match.start()-10):match.start()]:
                continue
            fake_val = get_fake_value(original, label)
            redacted_text = redacted_text[:match.start()] + fake_val + redacted_text[match.end():]
    return redacted_text

def redact_docx_document(doc):
    """Processes all paragraphs and table cells in a Document object in fast batches."""
    replacement_map.clear()
    
    # Collect all paragraph objects from document body and tables
    all_paras = list(doc.paragraphs)
    processed_cells = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell._tc in processed_cells:
                    continue
                processed_cells.add(cell._tc)
                all_paras.extend(cell.paragraphs)

    # Step 1: Regex redaction & collect unique texts needing NER
    para_regex_map = {}
    texts_for_ner = set()

    for para in all_paras:
        txt = para.text
        if not txt.strip():
            continue
        if txt not in para_regex_map:
            reg_txt = redact_text_regex(txt)
            para_regex_map[txt] = reg_txt
            if re.search(r'[a-zA-Z]', reg_txt):
                texts_for_ner.add(reg_txt)

    # Step 2: Batch NER with C-compiled nlp.pipe
    unique_ner_texts = list(texts_for_ner)
    ner_results = {}
    
    if unique_ner_texts:
        docs = nlp.pipe(unique_ner_texts, batch_size=128, disable=["tagger", "parser", "attribute_ruler", "lemmatizer"])
        for orig_text, spacy_doc in zip(unique_ner_texts, docs):
            pii_entities = [ent for ent in spacy_doc.ents if ent.label_ in ["PERSON", "ORG", "GPE", "LOC"]]
            pii_entities = sorted(pii_entities, key=lambda x: x.start_char, reverse=True)
            
            redacted_text = orig_text
            for ent in pii_entities:
                fake_val = get_fake_value(ent.text, ent.label_)
                redacted_text = redacted_text[:ent.start_char] + fake_val + redacted_text[ent.end_char:]
            ner_results[orig_text] = redacted_text

    # Step 3: Apply final redacted texts back to paragraphs
    final_cache = {}
    for para in all_paras:
        txt = para.text
        if not txt.strip():
            continue
        if txt not in final_cache:
            reg_txt = para_regex_map[txt]
            final_cache[txt] = ner_results.get(reg_txt, reg_txt)
        para.text = final_cache[txt]

def redact_document_stream(stream):
    """Reads a docx stream, redacts it in memory, and returns a BytesIO stream with the redacted document."""
    try:
        doc = Document(stream)
    except Exception as e:
        print(f"Error: Could not open document stream. Exact error: {e}")
        return None

    redact_docx_document(doc)

    output_stream = io.BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)
    return output_stream

def process_word_document(input_filename):
    """Reads a docx file, redacts text in paragraphs and tables, and modifies it in-place."""
    print(f"Opening '{input_filename}'...")
    try:
        doc = Document(input_filename)
    except Exception as e:
        print(f"Error: Could not open the file. Make sure '{input_filename}' exists.")
        print(f"Exact error: {e}")
        return False

    print("Scrubbing document for PII (this might take a moment for large files)...")
    redact_docx_document(doc)

    doc.save(input_filename)
    print(f"Success! Document redacted and saved as '{input_filename}'")
    return True

if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python3 redactor.py <input_file>")
        print("Example: python3 redactor.py 'Red Herring Prospectus.docx'")
        sys.exit(1)
    
    input_file = sys.argv[1]
    
    # Run the processor - modifies the file in-place
    success = process_word_document(input_file)
    sys.exit(0 if success else 1)